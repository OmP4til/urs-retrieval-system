
import io
import re
from typing import List, Dict, Any
import pdfplumber
from docx import Document
from lxml import etree
import zipfile
import openpyxl

OOXML_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# Note: The following imports have external dependencies that must be installed.
# For OCR (PDFs with images):
# - pytesseract: pip install pytesseract
# - pdf2image: pip install pdf2image
# - Google Tesseract: Must be installed on the system (e.g., via brew, apt-get)
# - Poppler: Required by pdf2image (e.g., via brew, apt-get)
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# For PDF annotations and comments:
# - PyMuPDF: pip install PyMuPDF
try:
    import fitz  # PyMuPDF
    PDF_ANNOTATIONS_AVAILABLE = True
except ImportError:
    PDF_ANNOTATIONS_AVAILABLE = False

# For semantic matching and intelligent text understanding:
# - sentence-transformers: pip install sentence-transformers
try:
    from sentence_transformers import SentenceTransformer
    import numpy as np
    from sklearn.metrics.pairwise import cosine_similarity
    SEMANTIC_MATCHING_AVAILABLE = True
except ImportError:
    SEMANTIC_MATCHING_AVAILABLE = False


def clean_extracted_text(text: str) -> str:
    """
    Clean extracted text by removing unwanted characters and normalizing whitespace.
    
    Args:
        text (str): Raw extracted text
        
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Remove backslashes and pipe symbols, but preserve word boundaries
    cleaned = text.replace("\\", " ").replace("|", " ")
    
    # Remove excessive whitespace and normalize
    cleaned = " ".join(cleaned.split())
    
    # Remove tabs and normalize spacing
    cleaned = cleaned.replace("\t", " ")
    
    return cleaned.strip()


class IntelligentTextMatcher:
    """
    Intelligent text matching using semantic embeddings to understand different terminologies.
    Can match concepts even when different words are used.
    """
    
    def __init__(self, model_name="all-MiniLM-L6-v2"):
        """Initialize the semantic matcher with a pre-trained model."""
        self.model = None
        self.model_name = model_name
        self._cache = {}  # Cache for embeddings
        self._terminology_mappings = self._load_domain_terminology()
        
        try:
            from sentence_transformers import SentenceTransformer
            import numpy as np
            from sklearn.metrics.pairwise import cosine_similarity
            self.model = SentenceTransformer(model_name)
            print(f"✅ Semantic matcher initialized with {model_name}")
        except ImportError as e:
            print(f"❌ Required packages not available: {e}")
            print("Install with: pip install sentence-transformers scikit-learn")
        except Exception as e:
            print(f"❌ Failed to load semantic model: {e}")
    
    def _load_domain_terminology(self) -> Dict[str, List[str]]:
        """Load domain-specific terminology mappings for URS documents."""
        return {
            # Certificate/Documentation terms
            "certificate": ["cert", "certification", "documentation", "proof", "validation", "verification"],
            "test_report": ["test report", "testing document", "validation report", "verification document", "test result", "test protocol"],
            "material": ["material", "substance", "component", "part", "element"],
            
            # Technical terms
            "metallic": ["metal", "metallic", "steel", "aluminum", "alloy"],
            "non_metallic": ["non-metal", "non-metallic", "plastic", "polymer", "ceramic", "rubber"],
            "contact": ["contact", "touching", "interface", "surface", "connection"],
            
            # System terms
            "system": ["system", "application", "software", "platform", "solution"],
            "access": ["access", "login", "authentication", "authorization", "entry"],
            "security": ["security", "protection", "safety", "secure", "protected"],
            "audit": ["audit", "log", "trail", "record", "tracking", "monitoring"],
            
            # Hardware/Output terms
            "printer": ["printer", "print", "printing", "output", "report generation", "document output"],
            "scope": ["scope", "responsibility", "coverage", "boundary", "domain"],
            "integration": ["integration", "interface", "connection", "linking", "combining"],
            
            # Process terms
            "provide": ["provide", "supply", "deliver", "give", "furnish", "submit"],
            "required": ["required", "needed", "necessary", "mandatory", "must have"],
            "validate": ["validate", "verify", "confirm", "check", "test", "ensure"],
            
            # Quality terms
            "compliance": ["compliance", "adherence", "conformity", "accordance", "standard"],
            "specification": ["specification", "spec", "requirement", "standard", "criteria"],
            "performance": ["performance", "operation", "functioning", "behavior", "execution"],
            
            # Format terms
            "readable": ["readable", "human readable", "legible", "clear", "understandable"],
            "format": ["format", "formatting", "layout", "structure", "presentation"],
            "demand": ["demand", "request", "on-demand", "when needed", "as required"]
        }
    
    def _get_embedding(self, text: str) -> np.ndarray:
        """Get embedding for text with caching."""
        if not self.model or not text:
            return np.array([])
        
        # Clean and normalize text
        text = clean_extracted_text(text).lower()
        
        if text in self._cache:
            return self._cache[text]
        
        try:
            embedding = self.model.encode([text])[0]
            self._cache[text] = embedding
            
            # Limit cache size
            if len(self._cache) > 1000:
                # Remove oldest entry
                oldest_key = next(iter(self._cache))
                del self._cache[oldest_key]
            
            return embedding
        except Exception as e:
            print(f"Error generating embedding for '{text}': {e}")
            return np.array([])
    
    def calculate_semantic_similarity(self, text1: str, text2: str) -> float:
        """Calculate semantic similarity between two texts."""
        if not self.model:
            # Fallback to simple keyword matching
            return self._keyword_similarity(text1, text2)
        
        emb1 = self._get_embedding(text1)
        emb2 = self._get_embedding(text2)
        
        if emb1.size == 0 or emb2.size == 0:
            return self._keyword_similarity(text1, text2)
        
        try:
            from sklearn.metrics.pairwise import cosine_similarity
            # Calculate cosine similarity
            similarity = cosine_similarity([emb1], [emb2])[0][0]
            return max(0.0, min(1.0, similarity))  # Clamp to 0-1 range
        except Exception as e:
            print(f"Error calculating similarity: {e}")
            return self._keyword_similarity(text1, text2)
    
    def _keyword_similarity(self, text1: str, text2: str) -> float:
        """Fallback keyword-based similarity with enhanced terminology understanding."""
        text1_clean = clean_extracted_text(text1).lower()
        text2_clean = clean_extracted_text(text2).lower()
        
        # Direct substring matching
        if text1_clean in text2_clean or text2_clean in text1_clean:
            return 0.8
        
        # Terminology mapping based similarity
        words1 = set(text1_clean.split())
        words2 = set(text2_clean.split())
        
        # Check for direct word overlap
        common_words = words1.intersection(words2)
        if common_words:
            overlap_ratio = len(common_words) / max(len(words1), len(words2))
            if overlap_ratio > 0.3:
                return overlap_ratio * 0.7  # Scale down keyword-only matches
        
        # Enhanced terminology mappings with scoring
        concept_matches = []
        for concept, synonyms in self._terminology_mappings.items():
            concept_in_1 = any(syn in text1_clean for syn in synonyms)
            concept_in_2 = any(syn in text2_clean for syn in synonyms)
            
            if concept_in_1 and concept_in_2:
                concept_matches.append(concept)
        
        # Score based on number of concept matches
        if concept_matches:
            base_score = 0.5
            bonus_score = min(0.3, len(concept_matches) * 0.1)  # Bonus for multiple concept matches
            return base_score + bonus_score
        
        # Special case handling for printer/print relationship
        printer_terms = ["printer", "print", "printing"]
        output_terms = ["output", "report", "document", "trail"]
        
        has_printer_1 = any(term in text1_clean for term in printer_terms)
        has_printer_2 = any(term in text2_clean for term in printer_terms)
        has_output_1 = any(term in text1_clean for term in output_terms)
        has_output_2 = any(term in text2_clean for term in output_terms)
        
        if (has_printer_1 and has_output_2) or (has_output_1 and has_printer_2):
            return 0.55  # Medium-high confidence for printer-output relationships
        
        return 0.0
    
    def find_best_semantic_matches(self, comment_text: str, requirement_texts: List[str], 
                                 threshold: float = 0.3) -> List[Dict[str, Any]]:
        """
        Find the best semantic matches for a comment against a list of requirements.
        
        Args:
            comment_text: The comment to match
            requirement_texts: List of requirement texts to match against
            threshold: Minimum similarity threshold
            
        Returns:
            List of matches with similarity scores and explanations
        """
        matches = []
        
        for req_text in requirement_texts:
            if not req_text or len(req_text.strip()) < 10:
                continue
            
            similarity = self.calculate_semantic_similarity(comment_text, req_text)
            
            if similarity >= threshold:
                # Generate explanation for the match
                explanation = self._generate_match_explanation(comment_text, req_text, similarity)
                
                matches.append({
                    "requirement_text": req_text,
                    "similarity_score": similarity,
                    "match_type": "semantic" if self.model else "keyword",
                    "explanation": explanation
                })
        
        # Sort by similarity score (descending)
        matches.sort(key=lambda x: x["similarity_score"], reverse=True)
        return matches
    
    def _generate_match_explanation(self, comment: str, requirement: str, score: float) -> str:
        """Generate human-readable explanation for why texts match."""
        comment_clean = clean_extracted_text(comment).lower()
        req_clean = clean_extracted_text(requirement).lower()
        
        # Find common concepts
        explanations = []
        
        # Check for terminology mappings
        for concept, synonyms in self._terminology_mappings.items():
            comment_has_concept = any(syn in comment_clean for syn in synonyms)
            req_has_concept = any(syn in req_clean for syn in synonyms)
            
            if comment_has_concept and req_has_concept:
                explanations.append(f"Both mention {concept}")
        
        # Check for direct word matches
        comment_words = set(comment_clean.split())
        req_words = set(req_clean.split())
        common_words = comment_words.intersection(req_words)
        
        if common_words:
            significant_words = [w for w in common_words if len(w) > 3]
            if significant_words:
                explanations.append(f"Shared keywords: {', '.join(significant_words[:3])}")
        
        if score > 0.7:
            confidence = "High"
        elif score > 0.5:
            confidence = "Medium"
        else:
            confidence = "Low"
        
        explanation_text = " | ".join(explanations) if explanations else "Semantic similarity detected"
        return f"{confidence} confidence ({score:.2f}): {explanation_text}"


# Global semantic matcher instance
_semantic_matcher = None

def get_semantic_matcher() -> IntelligentTextMatcher:
    """Get or create the global semantic matcher instance."""
    global _semantic_matcher
    if _semantic_matcher is None:
        _semantic_matcher = IntelligentTextMatcher()
    return _semantic_matcher

    
def detect_document_structure(doc) -> Dict[str, Any]:
    """Detect document structure including sections, headers, and index."""
    structure = {
        "sections": {},
        "index_detected": False,
        "header_patterns": [],
        "toc_pages": []
    }
    
    # Enhanced section patterns for URS documents
    section_patterns = [
        r'^\d+\.0\s+[A-Z][^.]*$',  # 1.0 INTRODUCTION, 2.0 SCOPE, etc.
        r'^\d+\.\d+\s+[A-Z][^.]*$',  # 1.1 Purpose, 2.1 Overview, etc.
        r'^\d+\s+[A-Z][A-Z\s]+$',  # 1 INTRODUCTION, 2 SCOPE (all caps)
        r'^[A-Z][A-Z\s]{10,}$',  # Long ALL CAPS HEADERS
        r'^\d+\.\d+\.\d+\s+[A-Z][^.]*$',  # 1.1.1 Sub-requirements
        r'^\d+\.\d+\.\d+\.\d+\s+[A-Z][^.]*$',  # Deep nesting like 7.2.1.1
    ]
    
    # Common URS section names and patterns
    urs_sections = [
        "approval", "introduction", "scope", "overview", "objective", "purpose", 
        "functional", "technical", "operational", "performance", "safety", 
        "environmental", "regulatory", "compliance", "validation", "qualification", 
        "installation", "training", "documentation", "maintenance", "support", 
        "requirements", "specifications", "process", "cleaning", "general requirement",
        "control", "system", "health", "abbreviations", "references", "definitions",
        "revision history", "vendor acceptance", "statutory regulations"
    ]
    
    # URS-specific table of contents detection
    toc_indicators = [
        "table of contents", "contents", "index", "table of content",
        "sr. no.", "sr.no.", "topics", "page no", "page no."
    ]
    
    current_section = "Introduction"
    section_counter = 0
    
    # Analyze paragraphs for structure
    for para_idx, para in enumerate(doc.paragraphs):
        text = clean_extracted_text(para.text)
        if not text:
            continue
            
        # Check if this looks like a section header
        is_header = False
        header_level = 0
        
        # Pattern-based detection
        for pattern in section_patterns:
            if re.match(pattern, text):
                is_header = True
                if '.' in text:
                    header_level = text.count('.')
                break
        
        # URS section name detection
        text_lower = text.lower()
        for section_name in urs_sections:
            if section_name in text_lower and len(text) < 100:
                is_header = True
                break
        
        # Formatting-based detection
        if para.runs and not is_header:
            first_run = para.runs[0]
            is_bold = first_run.bold
            is_larger = False
            try:
                font_size = first_run.font.size
                if font_size and font_size.pt > 12:
                    is_larger = True
            except:
                pass
            
            # Check if it's a numbered heading
            numbered_heading = re.match(r'^\d+(\.\d+)*\s+', text)
            
            if (is_bold or is_larger or numbered_heading) and len(text) < 100 and len(text) > 5:
                is_header = True
        
        if is_header:
            section_counter += 1
            
            # Clean up section name
            section_name = text
            # Remove numbering for cleaner section names
            section_name = re.sub(r'^\d+(\.\d+)*\s*', '', section_name)
            section_name = section_name.title()
            
            if not section_name:
                section_name = f"Section {section_counter}"
            
            current_section = section_name
            structure["sections"][section_name] = {
                "title": text,
                "paragraph_index": para_idx,
                "subsections": [],
                "requirements": [],
                "section_number": section_counter,
                "level": header_level
            }
            structure["header_patterns"].append(text)
            
        # Detect table of contents
        toc_keywords = ["table of contents", "contents", "index", "table of content"]
        if any(keyword in text.lower() for keyword in toc_keywords):
            structure["index_detected"] = True
            structure["toc_pages"].append(para_idx)
    
    return structure


def get_docx_comments_with_text_mapping(docx_file_bytes: bytes) -> Dict[str, Dict[str, Any]]:
    """Return dict mapping text content to comments for a DOCX file with precise sentence-level mapping."""
    try:
        with zipfile.ZipFile(io.BytesIO(docx_file_bytes)) as z:
            # Get comments
            if "word/comments.xml" not in z.namelist():
                return {}

            comments_xml = z.read("word/comments.xml")
            comments_et = etree.XML(comments_xml)

            # Parse all comments first
            comments_dict = {}
            comments = comments_et.xpath("//w:comment", namespaces=OOXML_NS)
            for c in comments:
                cid = c.xpath("@w:id", namespaces=OOXML_NS)[0]
                text = c.xpath("string(.)", namespaces=OOXML_NS)
                author = (c.xpath("@w:author", namespaces=OOXML_NS) or [None])[0]
                initials = (c.xpath("@w:initials", namespaces=OOXML_NS) or [None])[0]
                date = (c.xpath("@w:date", namespaces=OOXML_NS) or [None])[0]
                comments_dict[cid] = {
                    "id": cid,
                    "text": text,
                    "author": author,
                    "initials": initials,
                    "date": date,
                }

            # Parse document to map comments to EXACT text positions
            document_xml = z.read("word/document.xml")
            doc_root = etree.fromstring(document_xml)

            # Find all comment references and their PRECISE text context
            text_to_comments = {}
            comment_refs = doc_root.xpath('.//w:commentReference', namespaces=OOXML_NS)

            for ref in comment_refs:
                comment_id = ref.get(f"{{{OOXML_NS['w']}}}id")

                if comment_id not in comments_dict:
                    continue

                comment_mapped = False

                # Method 1: Get the EXACT text run that contains the comment reference
                parent_run = ref.xpath('ancestor::w:r', namespaces=OOXML_NS)
                
                if parent_run:
                    # Get the specific text content of this run
                    run_text_nodes = parent_run[0].xpath('.//w:t', namespaces=OOXML_NS)
                    run_text = clean_extracted_text(''.join([node.text for node in run_text_nodes if node.text]))
                    
                    if run_text and len(run_text) > 3:  # Only meaningful text
                        if run_text not in text_to_comments:
                            text_to_comments[run_text] = []
                        text_to_comments[run_text].append(comments_dict[comment_id])
                        comment_mapped = True
                        continue

                # Method 2: If in a table, try table cell context
                if not comment_mapped:
                    parent_cell = ref.xpath('ancestor::w:tc', namespaces=OOXML_NS)
                    if parent_cell:
                        # Get all text from the table cell
                        cell_text_nodes = parent_cell[0].xpath('.//w:t', namespaces=OOXML_NS)
                        cell_text = clean_extracted_text(''.join([node.text for node in cell_text_nodes if node.text]))
                        
                        if cell_text and len(cell_text) > 3:
                            # For table cells, be more selective about what we map to
                            # Split by common delimiters and map to the most relevant part
                            parts = re.split(r'[:\s]{2,}', cell_text)  # Split on colons and multiple spaces
                            
                            # Find the most relevant part (shortest meaningful text that might be a heading/key)
                            best_part = None
                            for part in parts:
                                part = clean_extracted_text(part)
                                if 3 < len(part) < 50:  # Prefer shorter, meaningful text
                                    if not best_part or len(part) < len(best_part):
                                        best_part = part
                            
                            if best_part:
                                if best_part not in text_to_comments:
                                    text_to_comments[best_part] = []
                                text_to_comments[best_part].append(comments_dict[comment_id])
                                comment_mapped = True
                            else:
                                # Fallback to full cell text if no good part found
                                if cell_text not in text_to_comments:
                                    text_to_comments[cell_text] = []
                                text_to_comments[cell_text].append(comments_dict[comment_id])
                                comment_mapped = True

                # Method 3: Enhanced paragraph-level mapping with better precision
                if not comment_mapped:
                    parent_para = ref.xpath('ancestor::w:p', namespaces=OOXML_NS)
                    
                    if parent_para:
                        # Get all text from the paragraph
                        para_text_nodes = parent_para[0].xpath('.//w:t', namespaces=OOXML_NS)
                        para_text = clean_extracted_text(''.join([node.text for node in para_text_nodes if node.text]))
                        
                        if para_text and len(para_text) > 3:
                            # Only map to paragraph if it's reasonably short (likely a heading or key phrase)
                            if len(para_text) < 100:
                                if para_text not in text_to_comments:
                                    text_to_comments[para_text] = []
                                text_to_comments[para_text].append(comments_dict[comment_id])
                                comment_mapped = True
                            else:
                                # For longer paragraphs, try to find the most relevant sentence
                                sentences = re.split(r'[.!?]+\s+', para_text)
                                
                                # Look for sentences that might be headings or key phrases
                                best_sentence = None
                                for sentence in sentences:
                                    sentence = clean_extracted_text(sentence)
                                    if 10 < len(sentence) < 80:  # Reasonable sentence length
                                        # Prefer sentences with keywords that might indicate sections
                                        if any(keyword in sentence.lower() for keyword in ['training', 'specification', 'requirement', 'shall', 'must']):
                                            best_sentence = sentence
                                            break
                                        elif not best_sentence:
                                            best_sentence = sentence
                                
                                if best_sentence:
                                    if best_sentence not in text_to_comments:
                                        text_to_comments[best_sentence] = []
                                    text_to_comments[best_sentence].append(comments_dict[comment_id])
                                    comment_mapped = True

            return text_to_comments

    except Exception as e:
        # Log the error for debugging but don't crash
        print(f"Warning: Error parsing DOCX comments: {e}")
        import traceback
        traceback.print_exc()
    return {}
def get_paragraph_comments(paragraph, comments_dict: Dict[str, Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Return list of comment dicts attached to a paragraph.

    Each item contains: {id, text, author, initials, date}
    """
    comments: List[Dict[str, Any]] = []
    try:
        # Get the XML element for this paragraph
        p_element = paragraph._element
        
        # Find all comment reference elements using lxml etree
        from lxml import etree
        # Convert paragraph element to lxml element if needed
        if hasattr(p_element, 'xml'):
            # This is a python-docx element, get raw XML
            xml_str = p_element.xml
            element = etree.fromstring(xml_str)
        else:
            element = p_element
            
        refs = element.xpath(".//w:commentReference", namespaces=OOXML_NS)
        for ref in refs:
            cid = ref.get(f"{{{OOXML_NS['w']}}}id")
            if cid and cid in comments_dict:
                comments.append(comments_dict[cid])
    except Exception:
        pass
    return comments


def categorize_requirements_by_section(pages: List[Dict[str, Any]], structure: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Categorize requirements by document sections."""
    
    # If no clear structure detected, return as-is
    if not structure["sections"]:
        return pages
    
    # Create section mapping
    section_mapping = {}
    sections_list = list(structure["sections"].items())
    
    for i, page in enumerate(pages):
        content = page.get("content", "")
        page_type = page.get("content_type", "")
        
        # Skip non-table content for now (focus on requirements in tables)
        if page_type != "docx-table-row":
            continue
            
        # Try to determine which section this requirement belongs to
        detected_section = None
        
        # Method 1: Look for section numbers in the content
        for section_title, section_info in sections_list:
            # Extract section number from title (e.g., "1.0", "2.1", "3.2")
            section_match = re.search(r'^(\d+(?:\.\d+)?)', section_title)
            if section_match:
                section_num = section_match.group(1)
                # Look for this section number in the table row content
                if section_num in content:
                    detected_section = section_title
                    break
        
        # Method 2: Look for keywords from section titles
        if not detected_section:
            for section_title, section_info in sections_list:
                # Extract key words from section title
                title_words = re.findall(r'\b[A-Z][a-z]+\b', section_title)
                if title_words:
                    for word in title_words:
                        if word.lower() in content.lower() and len(word) > 3:
                            detected_section = section_title
                            break
                if detected_section:
                    break
        
        # Method 3: Sequential assignment based on position
        if not detected_section and sections_list:
            # For table rows, try to assign based on order
            section_index = min(i // 10, len(sections_list) - 1)  # Rough grouping
            detected_section = sections_list[section_index][0]
        
        # Add section information to the page
        if detected_section:
            page["section"] = detected_section
            page["section_number"] = structure["sections"][detected_section].get("section_number", 0)
        else:
            page["section"] = "Unspecified"
            page["section_number"] = 999
    
    return pages


def extract_section_based_requirements(doc, file_bytes: bytes) -> List[Dict[str, Any]]:
    """Extract requirements organized by document sections."""
    
    # First, detect document structure
    structure = detect_document_structure(doc)
    text_to_comments = get_docx_comments_with_text_mapping(file_bytes)
    
    pages = []
    content_buffer = []
    comments_buffer = []
    page_num = 1
    current_section = "Introduction"
    
    # Create a lookup for section detection by keywords
    section_keywords = {}
    for section_name, section_info in structure["sections"].items():
        words = section_name.lower().split()
        for word in words:
            if len(word) > 3:  # Only meaningful words
                if word not in section_keywords:
                    section_keywords[word] = []
                section_keywords[word].append(section_name)
    
    # Iterate through document elements
    for element in doc.element.body:
        if element.tag.endswith('p'):
            # Find the corresponding paragraph object
            para = None
            for p in doc.paragraphs:
                if p._p == element:
                    para = p
                    break
            
            if para and para.text.strip():
                para_text = clean_extracted_text(para.text)
                
                # Enhanced section detection
                detected_section = None
                
                # Method 1: Direct section title match
                for section_name in structure["sections"]:
                    section_title = structure["sections"][section_name]["title"]
                    if section_title.lower() in para_text.lower() or para_text.lower() in section_title.lower():
                        detected_section = section_name
                        break
                
                # Method 2: Check for numbered sections in the paragraph
                if not detected_section:
                    numbered_match = re.search(r'(\d+\.?\d*)\s+([A-Z][A-Za-z\s]+)', para_text)
                    if numbered_match:
                        section_text = clean_extracted_text(numbered_match.group(2))
                        # Look for this section text in our detected sections
                        for section_name in structure["sections"]:
                            if section_text.lower() in section_name.lower():
                                detected_section = section_name
                                break
                
                # Method 3: Keyword-based detection
                if not detected_section:
                    para_words = para_text.lower().split()
                    for word in para_words:
                        if word in section_keywords:
                            # Use the first matching section
                            detected_section = section_keywords[word][0]
                            break
                
                if detected_section:
                    current_section = detected_section
                
                content_buffer.append(para_text)
                para_comments = text_to_comments.get(para_text, [])
                comments_buffer.extend(para_comments)
        
        elif element.tag.endswith('tbl'):
            # Find the corresponding table object
            table = None
            for t in doc.tables:
                if t._tbl == element:
                    table = t
                    break
            
            if not table:
                continue
            
            # Add any preceding text paragraphs as a page
            if content_buffer:
                pages.append({
                    "page_number": page_num,
                    "content": "\n".join(content_buffer).strip(),
                    "tables": [],
                    "content_type": "docx",
                    "comments": comments_buffer,
                    "section": current_section,
                })
                page_num += 1
                content_buffer, comments_buffer = [], []
            
            # Process table with enhanced section awareness
            table_header = None
            
            for row_idx, row in enumerate(table.rows):
                row_text_cells = []
                row_comments = []
                row_section = current_section  # Default to current section
                
                for cell in row.cells:
                    cell_text = clean_extracted_text(cell.text)
                    row_text_cells.append(cell_text)
                    
                    # Enhanced section detection for URS documents
                    section_number_match = re.search(r'(\d+\.\d+(\.\d+)?(\.\d+)?|\d+\.0|\d+)\s', cell_text)
                    if section_number_match:
                        section_num = section_number_match.group(1)
                        
                        # Map section numbers to meaningful sections based on URS structure
                        if section_num.startswith('1'):
                            if '1.0' in section_num:
                                row_section = "Approval"
                            else:
                                row_section = "Introduction"
                        elif section_num.startswith('2'):
                            row_section = "Introduction"
                        elif section_num.startswith('3'):
                            row_section = "Overview" 
                        elif section_num.startswith('4'):
                            row_section = "Statutory Regulations"
                        elif section_num.startswith('5'):
                            row_section = "Process Requirements"
                        elif section_num.startswith('6'):
                            row_section = "Cleaning Requirements"
                        elif section_num.startswith('7'):
                            row_section = "General Requirement Functions"
                        elif section_num.startswith('8'):
                            row_section = "Installation Requirements"
                        elif section_num.startswith('9'):
                            row_section = "Health Safety Environment"
                        elif section_num.startswith('10'):
                            row_section = "Documentation Training"
                        elif section_num.startswith('11'):
                            row_section = "References Definitions"
                        elif section_num.startswith('12'):
                            row_section = "Abbreviations"
                        elif section_num.startswith('13'):
                            row_section = "Revision History"
                        elif section_num.startswith('14'):
                            row_section = "Vendor Acceptance"
                        else:
                            row_section = f"Section {section_num}"
                    
                    # Also check for common URS keywords in the cell to refine section detection
                    cell_lower = cell_text.lower()
                    if not section_number_match:
                        if any(keyword in cell_lower for keyword in ["control", "operational", "parameter"]):
                            row_section = "General Requirement Functions"
                        elif any(keyword in cell_lower for keyword in ["process", "granulation", "flow"]):
                            row_section = "Process Requirements"
                        elif any(keyword in cell_lower for keyword in ["cleaning", "cip", "wash"]):
                            row_section = "Cleaning Requirements"
                        elif any(keyword in cell_lower for keyword in ["installation", "utility", "connection"]):
                            row_section = "Installation Requirements"
                        elif any(keyword in cell_lower for keyword in ["training", "documentation", "manual"]):
                            row_section = "Documentation Training"
                        elif any(keyword in cell_lower for keyword in ["safety", "health", "environment", "emergency"]):
                            row_section = "Health Safety Environment"
                    
                    # INTELLIGENT comment matching with semantic understanding
                    cell_comments = []
                    semantic_matcher = get_semantic_matcher()
                    
                    # Stage 1: EXACT text matching (highest priority)
                    if cell_text and cell_text in text_to_comments:
                        cell_comments.extend(text_to_comments[cell_text])
                    
                    # Stage 2: Sentence-level matching within cell
                    if cell_text and len(cell_text) > 20:
                        # Split cell content into sentences
                        sentences = re.split(r'[.!?]+\s+', cell_text)
                        for sentence in sentences:
                            sentence = clean_extracted_text(sentence)
                            if sentence and sentence in text_to_comments:
                                cell_comments.extend(text_to_comments[sentence])
                    
                    # Stage 3: Individual paragraph matching within cell
                    for para in cell.paragraphs:
                        para_text = clean_extracted_text(para.text)
                        if para_text and para_text in text_to_comments:
                            cell_comments.extend(text_to_comments[para_text])
                        
                        # Also split paragraph into sentences for precise mapping
                        if para_text and len(para_text) > 20:
                            para_sentences = re.split(r'[.!?]+\s+', para_text)
                            for sent in para_sentences:
                                sent = clean_extracted_text(sent)
                                if sent and sent in text_to_comments:
                                    cell_comments.extend(text_to_comments[sent])
                    
                    # Stage 4: Semantic matching for unmatched requirements (NEW!)
                    if not cell_comments and cell_text and len(cell_text) > 15:
                        # Get all available comment texts for semantic matching
                        all_comment_texts = []
                        for comment_list in text_to_comments.values():
                            for comment in comment_list:
                                comment_text = comment.get('text', '')
                                if comment_text and comment_text not in [c.get('text', '') for c in cell_comments]:
                                    all_comment_texts.append({
                                        'text': comment_text,
                                        'comment_obj': comment
                                    })
                        
                        # Find semantically similar comments
                        comment_text_list = [ct['text'] for ct in all_comment_texts]
                        if comment_text_list:
                            semantic_matches = semantic_matcher.find_best_semantic_matches(
                                cell_text, 
                                comment_text_list,
                                threshold=0.4  # Higher threshold for comment matching
                            )
                            
                            # Add the best semantic matches
                            for match in semantic_matches[:2]:  # Limit to top 2 semantic matches
                                # Find the corresponding comment object
                                for ct in all_comment_texts:
                                    if ct['text'] == match['requirement_text']:
                                        # Create enhanced comment with semantic info
                                        enhanced_comment = ct['comment_obj'].copy()
                                        enhanced_comment['semantic_match'] = True
                                        enhanced_comment['similarity_score'] = match['similarity_score']
                                        enhanced_comment['match_explanation'] = match['explanation']
                                        cell_comments.append(enhanced_comment)
                                        break
                    
                    # Add unique comments with safety limit
                    for comment in cell_comments:
                        comment_key = (comment.get('id'), comment.get('text', ''))
                        if not any((c.get('id'), c.get('text', '')) == comment_key for c in row_comments):
                            if len(row_comments) < 5:  # Limit for precision
                                row_comments.append(comment)
                
                row_content = " | ".join(row_text_cells)
                
                if not row_content.strip():
                    continue
                
                # Skip table of contents rows (they're not requirements)
                row_lower = row_content.lower()
                is_toc_row = (
                    # Very specific TOC patterns only
                    (len(row_content.split('|')) <= 3 and  # Simple 2-3 column structure
                     any(toc_indicator in row_lower for toc_indicator in [
                         "table of contents", "sr. no", "sr.no", "topics", "page no"
                     ]) and
                     # Must have page numbers or just section titles
                     (re.search(r'\b\d+(-\d+)?\s*$', row_content.strip()) or  # Ends with page numbers
                      re.search(r'^\d+\.0\s+[A-Z\s]+\s*\d*(-\d+)?$', row_content))) or
                    
                    # Header rows that are clearly just navigation
                    (row_lower in ["sr. no. | topics | page no.", "sr.no | topics | page no", 
                                   "no. | component/features | description"])
                )
                
                # Be much more selective about what we exclude
                if is_toc_row:
                    continue  # Skip only clear table of contents entries
                
                # Update current section for subsequent rows
                current_section = row_section
                
                # Enhanced header detection for URS tables
                is_header_row = False
                if row_idx == 0:
                    # Be more selective about what constitutes a header
                    header_indicators = [
                        "sr. no", "sr.no", "topics", "page no", "page no.",
                        "component", "features", "description"
                    ]
                    row_lower = row_content.lower()
                    
                    # Only exclude if it's clearly just a header with no content
                    is_header_row = (
                        not row_comments and  # Headers typically don't have comments
                        len(row_content.strip()) < 50 and  # Short header-like content
                        any(indicator in row_lower for indicator in header_indicators) and
                        not any(req_word in row_lower for req_word in ["shall", "must", "should", "required", "specification"])
                    )
                
                # Don't skip first rows that contain actual requirements
                if is_header_row and row_idx == 0:
                    table_header = row_content
                    continue
                
                # Create page for each requirement row
                row_display_content = row_content
                if table_header:
                    row_display_content = f"{table_header}\n{row_content}"
                
                if row_content.strip():
                    pages.append({
                        "page_number": f"Table-{page_num}-Row-{row_idx+1}",
                        "content": row_display_content,
                        "tables": [row_display_content],
                        "content_type": "docx-table-row",
                        "comments": row_comments,
                        "section": row_section,
                        "requirement_id": f"{row_section.replace(' ', '')}-{page_num}-{row_idx+1}",
                    })
                    page_num += 1
    
    # Add any remaining content
    if content_buffer:
        pages.append({
            "page_number": page_num,
            "content": "\n".join(content_buffer).strip(),
            "tables": [],
            "content_type": "docx",
            "comments": comments_buffer,
            "section": current_section,
        })
    
    return pages


def extract_pdf_annotations_and_text(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text and annotations from PDF with precise line-by-line mapping.
    Returns dict with text content and mapped annotations.
    """
    if not PDF_ANNOTATIONS_AVAILABLE:
        print("PyMuPDF not available, falling back to basic PDF extraction")
        return {"pages": [], "annotations": []}
    
    try:
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_data = []
        all_annotations = []
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            
            # Extract text with line-level precision
            text_dict = page.get_text("dict")
            page_text = ""
            text_lines = []
            line_positions = []
            
            # Build text line by line with position tracking
            for block in text_dict["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        line_text = ""
                        line_bbox = None
                        
                        for span in line["spans"]:
                            line_text += span["text"]
                            if line_bbox is None:
                                line_bbox = span["bbox"]
                            else:
                                # Expand bbox to include this span
                                line_bbox = [
                                    min(line_bbox[0], span["bbox"][0]),
                                    min(line_bbox[1], span["bbox"][1]),
                                    max(line_bbox[2], span["bbox"][2]),
                                    max(line_bbox[3], span["bbox"][3])
                                ]
                        
                        if line_text.strip():
                            cleaned_line_text = clean_extracted_text(line_text)
                            if cleaned_line_text:  # Only add if there's content after cleaning
                                text_lines.append(cleaned_line_text)
                                line_positions.append(line_bbox)
                                page_text += cleaned_line_text + "\n"
            
            # Extract annotations for this page
            page_annotations = []
            annotations = page.annots()
            
            for annot in annotations:
                try:
                    annot_dict = annot.info
                    annot_rect = annot.rect
                    
                    # Get annotation content
                    content = clean_extracted_text(annot_dict.get("content", ""))
                    if not content:
                        # Try to get subject or title as fallback
                        content = clean_extracted_text(annot_dict.get("subject", "") or annot_dict.get("title", ""))
                    
                    if content:
                        # Find which text line(s) this annotation is closest to
                        closest_lines = []
                        annot_center_y = (annot_rect.y0 + annot_rect.y1) / 2
                        
                        # Find overlapping or nearby text lines
                        for i, line_bbox in enumerate(line_positions):
                            if line_bbox:
                                line_center_y = (line_bbox[1] + line_bbox[3]) / 2
                                vertical_distance = abs(annot_center_y - line_center_y)
                                
                                # Consider lines within reasonable vertical proximity
                                if vertical_distance < 50:  # Adjust threshold as needed
                                    closest_lines.append({
                                        "line_index": i,
                                        "line_text": text_lines[i],
                                        "distance": vertical_distance
                                    })
                        
                        # Sort by distance and take the closest
                        closest_lines.sort(key=lambda x: x["distance"])
                        
                        annotation_data = {
                            "id": f"page_{page_num + 1}_annot_{len(page_annotations)}",
                            "type": annot_dict.get("name", "Comment"),
                            "text": content,
                            "author": annot_dict.get("author", "Unknown"),
                            "page": page_num + 1,
                            "position": {
                                "x0": annot_rect.x0,
                                "y0": annot_rect.y0,
                                "x1": annot_rect.x1,
                                "y1": annot_rect.y1
                            },
                            "associated_lines": closest_lines[:3]  # Keep top 3 closest lines
                        }
                        
                        page_annotations.append(annotation_data)
                        all_annotations.append(annotation_data)
                        
                except Exception as e:
                    print(f"Error processing annotation on page {page_num + 1}: {e}")
                    continue
            
            # Extract tables using fitz
            tables = []
            try:
                tabs = page.find_tables()
                for tab in tabs:
                    table_data = tab.extract()
                    if table_data:
                        table_text = []
                        for row in table_data:
                            row_text = " | ".join(str(cell) if cell else "" for cell in row)
                            table_text.append(row_text)
                        tables.append("\n".join(table_text))
            except Exception as e:
                print(f"Error extracting tables from page {page_num + 1}: {e}")
            
            page_data = {
                "page_number": page_num + 1,
                "content": page_text.strip(),
                "text_lines": text_lines,
                "line_positions": line_positions,
                "tables": tables,
                "annotations": page_annotations,
                "content_type": "pdf"
            }
            
            pages_data.append(page_data)
        
        pdf_doc.close()
        
        return {
            "pages": pages_data,
            "annotations": all_annotations
        }
        
    except Exception as e:
        print(f"Error extracting PDF annotations: {e}")
        return {"pages": [], "annotations": []}


def map_pdf_annotations_to_requirements(pages_data: List[Dict], annotations: List[Dict]) -> List[Dict]:
    """
    Map PDF annotations to specific requirements with intelligent semantic matching.
    Uses both spatial proximity and semantic understanding.
    """
    semantic_matcher = get_semantic_matcher()
    
    for page_data in pages_data:
        page_annotations = page_data.get("annotations", [])
        page_content = page_data.get("content", "")
        text_lines = page_data.get("text_lines", [])
        
        # Create a mapping from annotations to specific text segments
        mapped_comments = []
        
        for annotation in page_annotations:
            comment_text = annotation.get("text", "")
            associated_lines = annotation.get("associated_lines", [])
            
            # Stage 1: Spatial proximity matching (existing logic)
            spatial_matches = []
            for line_info in associated_lines:
                line_text = line_info.get("line_text", "")
                
                # Comprehensive URS requirement detection keywords (much more inclusive)
                requirement_keywords = [
                    # Obligation keywords (strong indicators)
                    "shall", "must", "should", "will", "require", "required", "necessary", "need", "needs", "needed",
                    "expect", "expected", "demand", "specify", "specified", "ensure", "provide", "maintain",
                    
                    # Technical specification keywords
                    "specification", "standard", "compliance", "certificate", "certification", "spec", "specs",
                    "test", "testing", "report", "validation", "verification", "qualification", "analysis", "analytical",
                    "measure", "measurement", "monitor", "monitoring", "control", "controlled", "verify",
                    
                    # Material and construction keywords
                    "material", "materials", "metallic", "non-metallic", "stainless steel", "aisi", "316l", "304l",
                    "product contact", "contact parts", "surface roughness", "mirror polished", "construction",
                    "built", "made", "fabricated", "manufactured", "design", "designed", "finish", "coating",
                    
                    # Process and operational keywords
                    "process", "processing", "operation", "operational", "control", "parameter", "parameters",
                    "temperature", "pressure", "flow", "speed", "capacity", "range", "limit", "limits",
                    "alarm", "trip", "setpoint", "operating", "performance", "efficiency", "function",
                    
                    # Documentation and training keywords
                    "documentation", "document", "training", "manual", "procedure", "procedures", "protocol",
                    "audit", "trail", "security", "access", "authorization", "backup", "record", "records",
                    "certificate", "certificates", "approval", "approved", "qualification", "qualified",
                    
                    # Equipment and system keywords
                    "equipment", "system", "systems", "machine", "device", "devices", "component", "components",
                    "assembly", "installation", "maintenance", "service", "support", "unit", "units", "module",
                    "instrument", "instruments", "tool", "tools", "apparatus", "machinery",
                    
                    # Quality and safety keywords
                    "quality", "safety", "gmp", "cgmp", "fda", "cfr", "part 11", "gamp", "iso", "astm",
                    "containment", "atex", "hazard", "hazards", "risk", "risks", "emergency", "protection",
                    "safe", "safer", "secure", "security", "reliable", "reliability",
                    
                    # Capacity and measurement keywords
                    "capacity", "volume", "size", "dimension", "dimensions", "area", "height", "width", "length",
                    "diameter", "weight", "mass", "density", "thickness", "level", "amount", "quantity",
                    
                    # Utility and infrastructure keywords
                    "utility", "utilities", "power", "electricity", "electrical", "water", "steam", "compressed air",
                    "nitrogen", "vacuum", "drainage", "ventilation", "hvac", "lighting", "communication",
                    
                    # Location and facility keywords
                    "location", "area", "room", "zone", "space", "facility", "building", "floor", "clean room",
                    "warehouse", "storage", "environment", "environmental", "atmosphere", "conditions",
                    
                    # Personnel and training keywords
                    "personnel", "staff", "operator", "operators", "technician", "manager", "supervisor",
                    "qualified", "trained", "competent", "authorized", "responsible", "user", "users",
                    
                    # Time and scheduling keywords
                    "time", "duration", "schedule", "frequency", "interval", "cycle", "batch", "continuous",
                    "periodic", "regular", "annual", "daily", "weekly", "monthly", "shift",
                    
                    # General action keywords
                    "provide", "supply", "deliver", "install", "configure", "setup", "implement", "execute",
                    "perform", "conduct", "carry out", "achieve", "accomplish", "complete", "finish",
                    
                    # Comparative and range keywords
                    "minimum", "maximum", "min", "max", "range", "between", "from", "to", "up to", "down to",
                    "less than", "greater than", "equal", "approximately", "about", "around", "typical",
                    
                    # Technical units and measurements
                    "kg", "g", "mg", "ton", "l", "ml", "m3", "m", "cm", "mm", "inch", "ft", "°c", "°f",
                    "bar", "psi", "mpa", "pascal", "rpm", "hz", "v", "volt", "amp", "watt", "kw", "mw",
                    "ph", "ppm", "ppb", "percent", "ratio", "conductivity", "viscosity",
                    
                    # Manufacturing and pharmaceutical specific
                    "batch", "lot", "campaign", "product", "intermediate", "api", "excipient", "tablet",
                    "capsule", "liquid", "solid", "powder", "granule", "coating", "formulation", "recipe"
                ]
                
                # URS-specific patterns for requirement identification
                line_lower = line_text.lower()
                line_stripped = line_text.strip()
                
                # Pattern 1: Numbered requirements (very strong indicator)
                numbered_requirement = (
                    re.search(r'^\d+\.\d+(\.\d+)?\s+', line_stripped) or  # 7.2.1.1, 10.2.3
                    re.search(r'^[•\-\*]\s*\d+\.\d+', line_stripped) or   # • 7.2.1
                    re.search(r'^\d+\)\s+', line_stripped) or             # 1) Item
                    re.search(r'^req\s*no\.?\s*\d+', line_lower) or       # Req No 7.1.1
                    re.search(r'^sr\.?\s*no\.?\s*\d+', line_lower)        # Sr. No. 1
                )
                
                # Pattern 2: Table-based requirements (common in URS) - much more inclusive
                table_requirement = (
                    '|' in line_text and len(line_text.split('|')) >= 2 and
                    len(line_stripped) > 15 and  # Any meaningful table content
                    not any(header in line_lower for header in ["sr. no", "sr.no", "page no", "topics"])
                )
                
                # Pattern 3: Specification lines with technical details - more inclusive
                specification_pattern = (
                    re.search(r'(temperature|pressure|flow|speed|capacity|voltage|frequency|ph|conductivity|viscosity).*[:=]\s*[\d\w]', line_lower) or
                    re.search(r'(min|max|minimum|maximum|range|limit|between|from|to)\.?\s*[:=]?\s*\d+', line_lower) or
                    re.search(r'\d+\s*(bar|°c|°f|rpm|m³/h|kg|g|mg|l|ml|mm|µm|inch|ft|v|amp|watt|hz|psi|mpa|ra\s*≤)', line_lower) or
                    re.search(r'(supply|operating|design|working|storage|ambient|process)\s+(pressure|temperature|conditions|range)', line_lower) or
                    re.search(r'\d+.*\s*(percent|%|ppm|ppb|ratio)', line_lower) or
                    re.search(r'(accuracy|precision|tolerance)\s*[:=±]\s*[\d.]+', line_lower)
                )
                
                # Pattern 4: Compliance and standard references - more inclusive
                compliance_pattern = (
                    re.search(r'(comply|compliance|conform|conformance|accordance|conformity)\s+with', line_lower) or
                    re.search(r'(21\s*cfr|part\s*11|gmp|cgmp|fda|eu|iso\s*\d+|astm|ansi|din|bs|en\s*\d+)', line_lower) or
                    re.search(r'(standard|standards|guideline|guidelines|regulation|regulations|directive|code|norm)', line_lower) or
                    re.search(r'(certificate|certification|qualified|approved|validated|verified)', line_lower)
                )
                
                # Pattern 5: Equipment specifications and requirements - much more inclusive
                equipment_pattern = (
                    re.search(r'(equipment|system|machine|device|unit|component|instrument|tool|apparatus)', line_lower) or
                    re.search(r'(material|construction|design|fabrication|manufacturing|installation)', line_lower) or
                    re.search(r'(provide|supply|deliver|install|configure|setup|implement|include)', line_lower) or
                    re.search(r'(capacity|volume|size|dimension|area|height|width|length|weight)', line_lower) or
                    re.search(r'(utility|utilities|power|electrical|water|steam|air|gas|vacuum|drainage)', line_lower)
                )
                
                # Pattern 6: Safety and operational requirements - more inclusive
                safety_pattern = (
                    re.search(r'(safety|emergency|alarm|interlock|guard|protection|containment|hazard|risk)', line_lower) or
                    re.search(r'(training|personnel|operator|qualified|authorized|competent)', line_lower) or
                    re.search(r'(maintenance|service|cleaning|calibration|repair|replacement)', line_lower) or
                    re.search(r'(documentation|record|report|manual|procedure|protocol)', line_lower)
                )
                
                # Pattern 7: Process and manufacturing requirements
                process_pattern = (
                    re.search(r'(process|processing|operation|batch|lot|campaign|production)', line_lower) or
                    re.search(r'(control|monitor|measure|test|analyze|verify|validate)', line_lower) or
                    re.search(r'(pharmaceutical|api|excipient|tablet|capsule|liquid|solid|powder)', line_lower)
                )
                
                # Pattern 8: Any line with technical specifications or numbers
                technical_pattern = (
                    re.search(r'\d+.*(?:kg|g|mg|l|ml|m|cm|mm|°c|°f|bar|psi|rpm|v|amp|watt|%|ppm)', line_lower) or
                    re.search(r'(?:range|between|from|to|up to|down to|approximately|about)\s*\d+', line_lower) or
                    re.search(r'\d+\s*(?:x|by|×)\s*\d+', line_lower)  # Dimensions like 100 x 200
                )
                
                # Combined requirement detection logic - much more inclusive
                is_requirement = (
                    len(line_text) > 8 and  # Reduced minimum length
                    (
                        # Strong indicators (high confidence)
                        numbered_requirement or
                        table_requirement or
                        specification_pattern or
                        compliance_pattern or
                        equipment_pattern or
                        safety_pattern or
                        process_pattern or
                        technical_pattern or
                        
                        # Keyword-based detection (medium confidence) - much more inclusive
                        (len(line_text) > 15 and 
                         any(keyword in line_lower for keyword in requirement_keywords) and
                         not any(exclusion in line_lower for exclusion in [
                             "table of contents", "page no.", "page no", "revision history", 
                             "abbreviation", "definition", "reference list", "approval signature",
                             "document control", "distribution list", "change control"
                         ])) or
                        
                        # Very inclusive catch-all for potential requirements
                        (len(line_text) > 20 and 
                         (
                             # Any line mentioning specifications, requirements, or technical details
                             any(term in line_lower for term in [
                                 "requirement", "specification", "standard", "certificate", "test", "report",
                                 "material", "equipment", "system", "process", "control", "parameter",
                                 "capacity", "temperature", "pressure", "flow", "safety", "training",
                                 "documentation", "maintenance", "utility", "design", "construction"
                             ]) or
                             
                             # Any line with technical measurements or ranges
                             re.search(r'\d+', line_text) and any(unit in line_lower for unit in [
                                 "kg", "g", "l", "ml", "m", "cm", "mm", "°c", "°f", "bar", "psi", "rpm", "%"
                             ]) or
                             
                             # Any line describing what should be provided/done
                             any(action in line_lower for action in [
                                 "provide", "supply", "deliver", "install", "ensure", "maintain", "control",
                                 "monitor", "verify", "validate", "comply", "conform", "include", "contain"
                             ])
                         ) and
                         not any(exclusion in line_lower for exclusion in [
                             "table of contents", "page no", "revision", "version", "date",
                             "author", "approved by", "reviewed by", "signature"
                         ]))
                    )
                )
                
                if is_requirement:
                    spatial_matches.append({
                        "text": line_text,
                        "match_type": "spatial",
                        "distance": line_info.get("distance", 0),
                        "keywords": [kw for kw in requirement_keywords if kw in line_lower]
                    })
            
            # Stage 2: Semantic matching for all potential requirements on the page - much more inclusive
            requirement_lines = [
                line for line in text_lines 
                if len(line) > 12 and (
                    # Strong requirement indicators
                    any(keyword in line.lower() for keyword in [
                        "shall", "must", "should", "will", "requirement", "required", "necessary",
                        "certificate", "test", "report", "material", "specification", "compliance",
                        "equipment", "system", "process", "control", "parameter", "capacity",
                        "temperature", "pressure", "flow", "safety", "training", "documentation",
                        "maintenance", "utility", "design", "construction", "provide", "supply",
                        "deliver", "install", "ensure", "maintain", "monitor", "verify", "validate"
                    ]) or
                    # Any line with technical measurements
                    (re.search(r'\d+', line) and any(unit in line.lower() for unit in [
                        "kg", "g", "mg", "l", "ml", "m3", "m", "cm", "mm", "°c", "°f", "bar", "psi", 
                        "rpm", "v", "amp", "watt", "hz", "%", "ppm", "ppb"
                    ])) or
                    # Table content (if contains |)
                    ('|' in line and len(line.split('|')) >= 2 and 
                     not any(header in line.lower() for header in ["sr. no", "page no", "topics"]))
                )
            ]
            
            semantic_matches = semantic_matcher.find_best_semantic_matches(
                comment_text, 
                requirement_lines, 
                threshold=0.3  # Lower threshold for semantic matching
            )
            
            # Stage 3: Combine and rank all matches
            all_matches = []
            
            # Add spatial matches with bonus score
            for spatial_match in spatial_matches:
                all_matches.append({
                    "requirement_text": spatial_match["text"],
                    "similarity_score": 0.9,  # High score for spatial proximity
                    "match_type": "spatial+keyword",
                    "explanation": f"Spatially close, Keywords: {', '.join(spatial_match['keywords'])}",
                    "distance": spatial_match["distance"]
                })
            
            # Add semantic matches
            for semantic_match in semantic_matches:
                # Check if this requirement was already matched spatially
                already_matched = any(
                    match["requirement_text"] == semantic_match["requirement_text"] 
                    for match in all_matches
                )
                
                if not already_matched:
                    all_matches.append(semantic_match)
                else:
                    # Enhance existing spatial match with semantic info
                    for match in all_matches:
                        if match["requirement_text"] == semantic_match["requirement_text"]:
                            match["similarity_score"] = min(1.0, match["similarity_score"] + 0.1)
                            match["match_type"] = "spatial+semantic"
                            match["explanation"] += f" | {semantic_match['explanation']}"
            
            # Stage 4: Select the best match(es)
            if all_matches:
                # Sort by similarity score and take the best match
                all_matches.sort(key=lambda x: x["similarity_score"], reverse=True)
                best_match = all_matches[0]
                
                # Only include matches above a reasonable threshold
                if best_match["similarity_score"] >= 0.3:
                    mapped_comments.append({
                        "id": annotation["id"],
                        "text": comment_text,
                        "author": annotation.get("author", "Unknown"),
                        "type": annotation.get("type", "Comment"),
                        "associated_text": best_match["requirement_text"],
                        "page": annotation["page"],
                        "precision": best_match["match_type"],
                        "similarity_score": best_match["similarity_score"],
                        "match_explanation": best_match["explanation"]
                    })
            
            # Stage 5: Special handling for unmatched high-value annotations
            if not all_matches and len(comment_text) > 10:
                # Try broader semantic search with lower threshold
                broader_matches = semantic_matcher.find_best_semantic_matches(
                    comment_text, 
                    text_lines,  # Search all lines, not just requirements
                    threshold=0.2
                )
                
                if broader_matches:
                    best_broad_match = broader_matches[0]
                    mapped_comments.append({
                        "id": annotation["id"],
                        "text": comment_text,
                        "author": annotation.get("author", "Unknown"),
                        "type": annotation.get("type", "Comment"),
                        "associated_text": best_broad_match["requirement_text"],
                        "page": annotation["page"],
                        "precision": "semantic-broad",
                        "similarity_score": best_broad_match["similarity_score"],
                        "match_explanation": f"Broad semantic match: {best_broad_match['explanation']}"
                    })
        
        # Update page data with mapped comments
        page_data["comments"] = mapped_comments
    
    return pages_data


def extract_structured_content(uploaded_file) -> List[Dict[str, Any]]:
    """
    Returns a list of page/chunk dicts for PDF/DOCX/XLSX/TXT.
    """
    try:
        file_bytes = uploaded_file.read()
        uploaded_file.seek(0)
    except Exception as e:
        print(f"Error reading file: {e}")
        return []

    name = getattr(uploaded_file, "name", "").lower()
    pages = []

    try:
        # ---------------- PDF ----------------
        if name.endswith(".pdf"):
            # Try enhanced annotation extraction first
            if PDF_ANNOTATIONS_AVAILABLE:
                print("Using enhanced PDF extraction with annotation support...")
                pdf_data = extract_pdf_annotations_and_text(file_bytes)
                enhanced_pages = map_pdf_annotations_to_requirements(
                    pdf_data["pages"], 
                    pdf_data["annotations"]
                )
                
                # Convert to the expected format and add section detection
                for page_data in enhanced_pages:
                    # Add OCR fallback if text is sparse
                    content = page_data.get("content", "")
                    if len(content.strip()) < 100 and OCR_AVAILABLE:
                        print(f"PDF page {page_data['page_number']}: Low text content, attempting OCR.")
                        try:
                            images = convert_from_bytes(
                                file_bytes, 
                                first_page=page_data['page_number'], 
                                last_page=page_data['page_number'], 
                                dpi=200
                            )
                            if images:
                                ocr_text = pytesseract.image_to_string(images[0]) or ""
                                content += "\n\n" + ocr_text
                                page_data["content"] = content
                        except Exception as ocr_error:
                            print(f"OCR failed for page {page_data['page_number']}: {ocr_error}")
                    
                    # Add tables to content if present
                    tables = page_data.get("tables", [])
                    if tables:
                        page_data["content"] += "\n\n--- TABLES ---\n" + "\n\n".join(tables)
                    
                    pages.append(page_data)
            
            else:
                # Fallback to basic extraction without annotations
                print("PyMuPDF not available, using basic PDF extraction...")
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for i, page in enumerate(pdf.pages, start=1):
                        # Improved text extraction to handle line breaks better
                        text = clean_extracted_text(page.extract_text(x_tolerance=2, y_tolerance=2) or "")
                        
                        # Fallback to OCR if text is sparse and dependencies are available
                        if len(text.strip()) < 100 and OCR_AVAILABLE:
                            print(f"PDF page {i}: Low text content, attempting OCR.")
                            try:
                                images = convert_from_bytes(
                                    file_bytes, first_page=i, last_page=i, dpi=200
                                )
                                if images:
                                    ocr_text = clean_extracted_text(pytesseract.image_to_string(images[0]) or "")
                                    text += "\n\n" + ocr_text # Append OCR text
                            except Exception as ocr_error:
                                print(f"OCR failed for page {i}: {ocr_error}")
                                print("Ensure Tesseract and Poppler are installed and in your PATH.")

                        # Extract tables separately
                        tables = []
                        try:
                            raw_tables = page.extract_tables()
                            for tbl in raw_tables:
                                rows = [clean_extracted_text(" | ".join(map(str, row))) for row in tbl]
                                cleaned_rows = [row for row in rows if row]  # Remove empty rows after cleaning
                                if cleaned_rows:
                                    tables.append("\n".join(cleaned_rows))
                        except Exception:
                            pass
                        
                        content = text
                        if tables:
                            content += "\n\n--- TABLES ---\n" + "\n\n".join(tables)

                        pages.append({
                            "page_number": i,
                            "content": content.strip(),
                            "tables": tables,
                            "content_type": "pdf",
                            "comments": [],
                        })

        # ---------------- DOCX ----------------
        elif name.endswith(".docx"):
            doc = Document(io.BytesIO(file_bytes))
            
            # Use section-based extraction for better organization
            pages = extract_section_based_requirements(doc, file_bytes)

        # ---------------- XLSX ----------------
        elif name.endswith((".xlsx", ".xls")):
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            for sheetname in wb.sheetnames:
                ws = wb[sheetname]
                rows_text = [clean_extracted_text(" | ".join(str(c) if c is not None else "" for c in row)) for row in ws.iter_rows(values_only=True)]
                cleaned_rows = [row for row in rows_text if row]  # Remove empty rows after cleaning
                content = "\n".join(cleaned_rows)
                if content.strip():
                    pages.append({
                        "page_number": sheetname,
                        "content": content.strip(),
                        "tables": [content],
                        "content_type": "xlsx",
                        "comments": [],
                    })

        # ---------------- TXT ----------------
        elif name.endswith(".txt"):
            try:
                text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                text = file_bytes.decode("latin-1", errors="ignore")
            pages.append({
                "page_number": 1,
                "content": text.strip(),
                "tables": [],
                "content_type": "txt",
                "comments": [],
            })

    except Exception as e:
        print(f"[extract_structured_content] error extracting '{name}': {e}")
        return []

    return pages
