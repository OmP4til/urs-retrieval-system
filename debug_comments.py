#!/usr/bin/env python3
"""
Debug comment extraction - test directly with DOCX files
"""

import os
import zipfile
import io
from lxml import etree
from docx import Document

OOXML_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

def debug_docx_comments(file_path):
    """Debug DOCX comment extraction step by step"""
    
    print(f"=== Debugging {file_path} ===")
    
    if not os.path.exists(file_path):
        print(f"File {file_path} not found!")
        return
    
    # Step 1: Check if comments.xml exists
    try:
        with zipfile.ZipFile(file_path, 'r') as docx_zip:
            file_list = docx_zip.namelist()
            print(f"Files in DOCX: {len(file_list)}")
            
            comments_files = [f for f in file_list if 'comment' in f.lower()]
            revisions_files = [f for f in file_list if any(x in f.lower() for x in ['revision', 'track', 'change'])]
            print(f"Comment-related files: {comments_files}")
            print(f"Revision-related files: {revisions_files}")
            print(f"All files: {file_list}")
            
            if 'word/comments.xml' not in file_list:
                print("❌ No word/comments.xml found - document has no comments")
                return
            else:
                print("✅ word/comments.xml found!")
                
            # Step 2: Read and parse comments.xml
            comments_xml = docx_zip.read('word/comments.xml')
            print(f"Comments XML size: {len(comments_xml)} bytes")
            
            # Parse XML
            et = etree.XML(comments_xml)
            print(f"XML root tag: {et.tag}")
            
            # Find all comments
            comments = et.xpath('//w:comment', namespaces=OOXML_NS)
            print(f"Found {len(comments)} comment elements")
            
            for i, c in enumerate(comments):
                cid = c.xpath('@w:id', namespaces=OOXML_NS)
                text = c.xpath('string(.)', namespaces=OOXML_NS)
                author = c.xpath('@w:author', namespaces=OOXML_NS)
                initials = c.xpath('@w:initials', namespaces=OOXML_NS)
                date = c.xpath('@w:date', namespaces=OOXML_NS)
                
                print(f"\nComment {i+1}:")
                print(f"  ID: {cid}")
                print(f"  Text: {text[:100]}...")
                print(f"  Author: {author}")
                print(f"  Initials: {initials}")
                print(f"  Date: {date}")
    
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return
    
    # Step 3: Test with python-docx
    try:
        print(f"\n=== Testing with python-docx ===")
        doc = Document(file_path)
        
        # Check paragraphs for comment references
        total_paras = len(doc.paragraphs)
        print(f"Document has {total_paras} paragraphs")
        
        paras_with_comments = 0
        for i, para in enumerate(doc.paragraphs):
            # Check each run for comment references
            for run in para.runs:
                try:
                    refs = run._r.xpath("./w:commentReference", namespaces=OOXML_NS)
                    if refs:
                        paras_with_comments += 1
                        print(f"Paragraph {i+1} has comment reference: {refs}")
                        break
                except Exception as e:
                    print(f"Error checking paragraph {i+1}: {e}")
        
        print(f"Paragraphs with comment references: {paras_with_comments}")
        
    except Exception as e:
        print(f"Error with python-docx: {e}")

def test_all_docx_files():
    """Test all DOCX files in current directory"""
    docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]
    
    if not docx_files:
        print("No DOCX files found in current directory")
        print("Available files:")
        for f in os.listdir('.'):
            if f.endswith(('.docx', '.doc')):
                print(f"  {f}")
        return
    
    for docx_file in docx_files:
        debug_docx_comments(docx_file)
        print("\n" + "="*50 + "\n")

if __name__ == "__main__":
    # Test specific file if provided
    test_files = [
        "URS Coating Machine Rev 1 - GLATT comments 03092025.docx",
        "Novugen_URS IGL (1).docx",
        "Novugen_URS IGL (1)",
        "test_urs_for_comments.docx"
    ]
    
    found_file = None
    for test_file in test_files:
        if os.path.exists(test_file):
            found_file = test_file
            break
        # Try with .docx extension if not present
        if not test_file.endswith('.docx'):
            test_file_with_ext = test_file + '.docx'
            if os.path.exists(test_file_with_ext):
                found_file = test_file_with_ext
                break
    
    if found_file:
        print(f"Testing file: {found_file}")
        debug_docx_comments(found_file)
    else:
        print(f"None of the test files found. Testing all DOCX files...")
        test_all_docx_files()