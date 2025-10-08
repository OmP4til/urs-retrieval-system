"""
Create a test DOCX file with comments to demonstrate comment extraction
"""

from docx import Document
from docx.oxml import OxmlElement, ns
import datetime

def create_test_docx_with_comments():
    # Create a new document
    doc = Document()
    
    # Add title
    doc.add_heading('Test URS with Comments', 0)
    
    # Add some requirements with simulated comments
    requirements = [
        "The system shall maintain temperature between 18-24°C at all times.",
        "The system must provide automated backup in case of power failure.", 
        "User interface should be intuitive and require minimal training.",
        "The system shall log all operations for audit purposes.",
        "Maintenance procedures must be documented and easily accessible."
    ]
    
    for i, req in enumerate(requirements):
        p = doc.add_paragraph(f"REQ-{i+1}: {req}")
        
        # Note: Adding comments programmatically to DOCX is complex
        # The comments will need to be added manually in Word
        # This script just creates the base document
    
    # Add instruction
    doc.add_paragraph("\nNOTE: To test comment extraction, please:")
    doc.add_paragraph("1. Open this file in Microsoft Word")
    doc.add_paragraph("2. Select some text in the requirements")
    doc.add_paragraph("3. Add comments using Review → New Comment")
    doc.add_paragraph("4. Include author information")
    doc.add_paragraph("5. Save the file")
    doc.add_paragraph("6. Upload it to the URS system")
    
    # Save the document
    doc.save('test_urs_for_comments.docx')
    print("Created test_urs_for_comments.docx")
    print("Please add comments in Word and then test the system!")

if __name__ == "__main__":
    create_test_docx_with_comments()